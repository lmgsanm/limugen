#!/usr/bin/env python3
# coding=gbk
__Author__="limugen"
import urllib.request
import urllib
import pymysql.cursors
import sys
import time
import os
import json
import datetime
import docx
from pyzabbix import ZabbixAPI
from docx.shared import Inches
from bs4 import BeautifulSoup
django_headers = {"User-Agent" : "Mozilla/5.0 (X11; U; Linux i686) Gecko/20071127 Firefox/2.0.0.11",
            "Content-Type" : "application/json",
            "Charset" : "utf8"
            }

def get_json_data(url):
    headers = {"User-Agent" : "Mozilla/5.0 (X11; U; Linux i686) Gecko/20071127 Firefox/2.0.0.11",
                "Content-Type" : "application/json",
                "Charset" : "utf8"
                }
    req = urllib.request.Request(url=url,headers=headers)
    req_result = urllib.request.urlopen(req).read().decode('utf-8')
    req_data = json.loads(req_result)
    return req_data

def get_group_info(group):
    host_dict = {}
    api = "http://10.205.56.124/api"
    api_url = '/'.join(["gethostbygroup",'='.join(["?groupname",'&'.join([])])])
    parament_last = '&'.join(['='.join(["?groupname",group]),"grouptype=projectteam"])
    parament_url = '/'.join(["gethostbygroup",parament_last])
    api_url = '/'.join([api,parament_url])
    host_info = []
    json_data = get_json_data(api_url)
    return json_data

def get_mysql_data(sql):
    host = "10.205.34.251"
    user = "readonly"
    password = "!@1234read"
    db = "zabbix"
    charset='utf8mb4'
    connection = pymysql.connect(host=host,user=user,password=password,db=db,charset=charset,cursorclass=pymysql.cursors.DictCursor)
    try:
        with connection.cursor() as cursor:
            sql = sql
            cursor.execute(sql)
            result = cursor.fetchall()
            return result
    finally:
        connection.close()

def get_host_info(group_name):
    group_list = get_group_info(group_name)
    host_info_dict = {}
    for host in group_list:
        ip = host["ipaddr"]
        hostname = host["hostname"]
        cpu = host["cpu"]
        memory = host["memory"]
        disk = host["disk"]
        host_api_url = '='.join(["http://10.205.56.124/api/getgroupbyhost/?ipaddr",str(ip)])
        addr_json = get_json_data(host_api_url)
        module_name = addr_json[0]["groupname"]
        host_info_list = [module_name,hostname,cpu,memory,disk]
        host_info_dict[ip] = host_info_list
    return host_info_dict

def get_zabbix_data(item_id):
    zabbix_server = "http://10.205.56.119"
    zabbix_user = "monitor"
    zabbix_password = "WSX@abc321,"
    item_id = item_id
    zabbix_api = ZabbixAPI(zabbix_server)
    zabbix_api.login(zabbix_user,zabbix_password)
    time_till = time.mktime(datetime.datetime.now().timetuple())
    time_from = time_till - 60 * 60 * 24 * 7
    history = zabbix_api.history.get(itemids=[item_id],
            time_till = time_till,
            time_from = time_from,
            output = 'extend',
            limit = '5000',
            )
    if not len(history):
        history = zabbix_api.history.get(itemids=[item_id],
                time_till = time_till,
                time_from = time_from,
                output = 'extend',
                limit = '5000',
                history = 0
                )
    value_temp = []
    for point in history:
        value_temp.append(int(point["value"]))
    num = len(value_temp)
    sum = 0
    for i in value_temp:
        sum = sum + float(i)
    value_temp_max = float(max(value_temp))
    value_temp_average = float(round(float(sum)/float(num),2))
    value_list = [value_temp_average,value_temp_max]
    return value_list

def get_host_monitor_data(host_info_dict,monitor_keys):
    host_monitor_data_dict = {}
    for ip in host_info_dict:
        zabbix_data_list = []
        for monitor_key in monitor_keys:
            item_id = get_item_data(ip,monitor_key)
            itemid = item_id[0].get("itemid")
            zabbix_data = get_zabbix_data(itemid)
            zabbix_data_list.extend(zabbix_data)
        host_monitor_data_dict[ip] = zabbix_data_list
    return host_monitor_data_dict

def get_item_data(ip,zabbix_key):
    sql = "select itemid from items where hostid in ( select hostid from hosts where host = \'" + ip + "\' ) AND key_ = \'" + zabbix_key + "\'"
    item_id = get_mysql_data(sql)
    return item_id

def get_graph_data(item_id):
    zabbix_server = "http://10.205.56.119"
    zabbix_user = "monitor"
    zabbix_password = "WSX@abc321,"
    item_id = item_id
    zabbix_api = ZabbixAPI(zabbix_server)
    zabbix_api.login(zabbix_user,zabbix_password)
    graph = zabbix_api.graph.get(itemids=[item_id])
    graph_list = [graph[0].get('graphid')]
    return graph_list

def get_host_graph_data(host_info_dict,monitor_keys):
    host_monitor_graph_dict = {}
    for ip in host_info_dict:
        zabbix_data_list = []
        for monitor_key in monitor_keys:
            item_id = get_item_data(ip,monitor_key)
            itemid = item_id[0].get("itemid")
            zabbix_data = get_graph_data(itemid)
            zabbix_data_list.append(zabbix_data)
        host_monitor_graph_dict[ip] = zabbix_data_list
    return host_monitor_graph_dict

def get_host_graph_url(graphid,period,stime):
    zabbix_graph_url = 'http://10.205.56.119/chart2.php'
    zabbix_graph__request_url = zabbix_graph_url + "?graphid=" + graphid + "&period=" + period + "&isNow=0&stime=" + stime
    return zabbix_graph__request_url

if __name__ == "__main__":
    hostgroup = sys.argv[1]
    monitor_keys = [ 'cpu.usage','vm.memory.usage','tcp.status["CLOSE_WAIT"]','tcp.status["ESTABLISHED"]','tcp.status["TIME_WAIT"]']
    host_info_data_dict = get_host_info(hostgroup)
    host_monitor_data_dict = get_host_monitor_data(host_info_data_dict,monitor_keys)
    host_monitor_graph_dict = get_host_graph_data(host_info_data_dict,monitor_keys)
    date = datetime.datetime.now().strftime('%Y-%m-%d')
    docx_name = sys.argv[1] + "系统运行周报" + "-" + date + ".docx"
    doc = docx.Document()
    doc_head_list = ["系统稳定率","系统故障/异常","系统变更","服务器监控","服务器监控图","数据库慢SQL","系统运行总结"]
#add 系统稳定率
    system_stability_head_list = ['系统名称','系统简码','异常时长','系统稳定率']
    doc.add_heading(doc_head_list[0],level=2)
    system_stability_table = doc.add_table(rows=1, cols=4, style='Table Grid')
    system_stability_head_cells = system_stability_table.rows[0].cells
    system_stability_head_cells[0].text = system_stability_head_list[0]
    system_stability_head_cells[1].text = system_stability_head_list[1]
    system_stability_head_cells[2].text = system_stability_head_list[2]
    system_stability_head_cells[3].text = system_stability_head_list[3]

#add 系统故障/异常
    system_abnormal_head_list = ['系统名称','系统简码','发生时间','持续时间','影响范围','故障/异常原因','解决方案','后续跟踪']
    doc.add_heading(doc_head_list[1],level=2)
    system_abnormal_table = doc.add_table(rows=1, cols=4, style='Table Grid')
    system_abnormal_head_cells = system_abnormal_table.rows[0].cells
    system_abnormal_head_cells[0].text = system_abnormal_head_list[0]
    system_abnormal_head_cells[1].text = system_abnormal_head_list[1]
    system_abnormal_head_cells[2].text = system_abnormal_head_list[2]
    system_abnormal_head_cells[3].text = system_abnormal_head_list[3]

#add 系统变更记
    system_alteration_head_list = ['版本发布','应用参数变更','服务器变更']
    system_version_alteration_head_list = ['版本发布时间','上线功能','备注']
    system_app_parameter_alteration_head_list = ['序号','模块','配置文件','配置项','配置参数值']
    system_machine_parameter_alteration_head_list = ['序号','ip地址','模块','变更内容']
   
    doc.add_heading(doc_head_list[2],level=2)
   
    doc.add_heading(system_alteration_head_list[0],level=3)
    system_version_alteration_table = doc.add_table(rows=1, cols=3, style='Table Grid')
    system_version_alteration_head_cells = system_version_alteration_table.rows[0].cells
    system_version_alteration_head_cells[0].text = system_version_alteration_head_list[0]
    system_version_alteration_head_cells[1].text = system_version_alteration_head_list[1]
    system_version_alteration_head_cells[2].text = system_version_alteration_head_list[2]

    doc.add_heading(system_alteration_head_list[1],level=3)
    system_app_parameter_alteration_table = doc.add_table(rows=1, cols=5, style='Table Grid')
    system_app_parameter_alteration_head_cells = system_app_parameter_alteration_table.rows[0].cells
    system_app_parameter_alteration_head_cells[0].text = system_app_parameter_alteration_head_list[0]
    system_app_parameter_alteration_head_cells[1].text = system_app_parameter_alteration_head_list[1]
    system_app_parameter_alteration_head_cells[2].text = system_app_parameter_alteration_head_list[2]
    system_app_parameter_alteration_head_cells[3].text = system_app_parameter_alteration_head_list[3]
    system_app_parameter_alteration_head_cells[4].text = system_app_parameter_alteration_head_list[4]
    
    doc.add_heading(system_alteration_head_list[2],level=3)
    system_machine_parameter_alteration_table = doc.add_table(rows=1, cols=4, style='Table Grid')
    system_machine_parameter_alteration_head_cells = system_machine_parameter_alteration_table.rows[0].cells
    system_machine_parameter_alteration_head_cells[0].text = system_machine_parameter_alteration_head_list[0]
    system_machine_parameter_alteration_head_cells[1].text = system_machine_parameter_alteration_head_list[1]
    system_machine_parameter_alteration_head_cells[2].text = system_machine_parameter_alteration_head_list[2]
    system_machine_parameter_alteration_head_cells[3].text = system_machine_parameter_alteration_head_list[3]


#add 服务器监控
    host_info_head_list = ["序号","IP地址","模块名称","主机名","CPU","内存","磁盘"]
    host_monitor_head_list = ["序号","IP地址","CPU平均使用率","CPU最大使用率","内存平均使用率","内存最大使用率","TCP-CLOSE_WAIT平均连接数","TCP-CLOSE_WAIT最大连接数","TCP-ESTABLISHED平均连接数","TCP-ESTABLISHED最大连接数","TCP-TIME_WAIT平均连接数","TCP-TIME_WAIT最大连接数"]
    doc.add_heading(doc_head_list[3],level=2)
#host_info_data_table
    doc.add_heading('服务器信息表',level=3)
    host_info_data_list = []
    for seq,ip in enumerate(host_info_data_dict,start=1):
        host_info_data_list.append([seq,ip,host_info_data_dict.get(ip)])
    host_info_rows = len(host_info_data_dict) + int(1)
    host_info_columns = len(host_info_head_list)
    host_info_data_table = doc.add_table(rows=1, cols=host_info_columns, style='Table Grid')
    host_info_head_cells = host_info_data_table.rows[0].cells
    host_info_head_cells[0].text = host_info_head_list[0]
    host_info_head_cells[1].text = host_info_head_list[1]
    host_info_head_cells[2].text = host_info_head_list[2]
    host_info_head_cells[3].text = host_info_head_list[3]
    host_info_head_cells[4].text = host_info_head_list[4]
    host_info_head_cells[5].text = host_info_head_list[5]
    host_info_head_cells[6].text = host_info_head_list[6]
    for item in host_info_data_list:
        host_info_data_rows_cells = host_info_data_table.add_row().cells
        host_info_data_rows_cells[0].text = str(item[0])
        host_info_data_rows_cells[1].text = str(item[1])
        host_info_data_rows_cells[2].text = str(item[2][0])
        host_info_data_rows_cells[3].text = str(item[2][1])
        host_info_data_rows_cells[4].text = str(item[2][2])
        host_info_data_rows_cells[5].text = str(item[2][3])
        host_info_data_rows_cells[6].text = str(item[2][4])

#host_monitor_data_table
    doc.add_heading('服务器监控信息',level=3)
    host_monitor_data_list = []
    for seq,ip in enumerate(host_monitor_data_dict,start=1):
        host_monitor_data_list.append([seq,ip,host_monitor_data_dict.get(ip)])
    host_monitor_rows = len(host_monitor_data_dict) + int(1)
    host_monitor_columns = len(host_monitor_head_list)
    host_monitor_data_table = doc.add_table(rows=1, cols=host_monitor_columns, style='Table Grid')
    host_monitor_head_cells = host_monitor_data_table.rows[0].cells
    host_monitor_head_cells[0].text = host_monitor_head_list[0]
    host_monitor_head_cells[1].text = host_monitor_head_list[1]
    host_monitor_head_cells[2].text = host_monitor_head_list[2]
    host_monitor_head_cells[3].text = host_monitor_head_list[3]
    host_monitor_head_cells[4].text = host_monitor_head_list[4]
    host_monitor_head_cells[5].text = host_monitor_head_list[5]
    host_monitor_head_cells[6].text = host_monitor_head_list[6]
    host_monitor_head_cells[7].text = host_monitor_head_list[7]
    host_monitor_head_cells[8].text = host_monitor_head_list[8]
    host_monitor_head_cells[9].text = host_monitor_head_list[9]
    host_monitor_head_cells[10].text = host_monitor_head_list[10]
    host_monitor_head_cells[11].text = host_monitor_head_list[11]
    #for item in host_monitor_data_dict:
    for item in host_monitor_data_list:
        host_monitor_data_rows_cells = host_monitor_data_table.add_row().cells
#        print(item)
        host_monitor_data_rows_cells[0].text = str(item[0])
        host_monitor_data_rows_cells[1].text = str(item[1])
        host_monitor_data_rows_cells[2].text = str(item[2][0])
        host_monitor_data_rows_cells[3].text = str(item[2][1])
        host_monitor_data_rows_cells[4].text = str(item[2][2])
        host_monitor_data_rows_cells[5].text = str(item[2][3])
        host_monitor_data_rows_cells[6].text = str(item[2][4])
        host_monitor_data_rows_cells[7].text = str(item[2][5])
        host_monitor_data_rows_cells[8].text = str(item[2][6])
        host_monitor_data_rows_cells[9].text = str(item[2][7])
        host_monitor_data_rows_cells[10].text = str(item[2][8])
        host_monitor_data_rows_cells[11].text = str(item[2][9])
   

#add 服务器监控图
    doc.add_heading(doc_head_list[4],level=2)
    period = str(60 * 60 * 24 * 7)
    starttime = str(datetime.datetime.now().strftime('%Y%m%d%H%M%S'))
    for ip in host_info_data_dict:
        head_name = str(ip) + "_" + str(host_info_data_dict.get(ip)[0])
        doc.add_heading(head_name,level=3)
        graphids = get_host_graph_data(host_info_data_dict,monitor_keys).get(ip)
        graphid_list = []
        for graphid in graphids:
            if graphid not in graphid_list:
                graphid_list.append(graphid)
        for graphid_id in graphid_list:
            graphid = str(graphid_id[0])
#            print(graphid)
#            sys.exit()
            graph_url = get_host_graph_url(graphid,period,starttime)
            graph_name = str(ip) + '_' + graphid + ".png"
            urllib.request.urlretrieve(graph_url,graph_name)
            doc.add_picture(graph_name,width=Inches(6))

#add 数据库慢SQL
    doc.add_heading(doc_head_list[5],level=2)

#add 系统运行总结
    doc.add_heading(doc_head_list[6],level=2)
    doc.save(docx_name)
